# =============================================================================
# modules/word_writer.py
# =============================================================================

import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from config import REPORT_LINGUA
from modules.ai_client import chiedi_ai


def _commento_ai(sezione: str, dati_json: str, kpi: dict) -> str:
    nav = kpi.get("nav", 0)
    prompt = f"""Sei un analista finanziario senior che redige una relazione professionale
    in {REPORT_LINGUA} per un cliente istituzionale.

    Scrivi 4-5 frasi di commento per la sezione "{sezione}".
    Book Value totale portafoglio: {nav:,.2f} €

    Dati aggregati:
    {dati_json}

    Linee guida:
    - Tono professionale e oggettivo
    - Evidenzia le voci più rilevanti per peso o variazione
    - Per il confronto Book Value - NAV nell'executive summary, utilizza l'analisi Book Value vs Fair Value (Paragrafo 2.9)
    - Segnala eventuali concentrazioni o elementi degni di nota
    - No elenchi puntati, solo prosa fluente
    - NON includere titoli o intestazioni nel testo
    - Inizia direttamente con il commento
    - Massimo 150 parole"""

    risposta = chiedi_ai(prompt, max_tokens=400)
    if risposta:
        import re
        risposta = re.sub(r'^#+\s*', '', risposta, flags=re.MULTILINE)
        risposta = re.sub(r'\*\*(.*?)\*\*', r'\1', risposta)
        risposta = re.sub(r'\*(.*?)\*', r'\1', risposta)
        risposta = risposta.strip()
    else:
        print(f"[AVVISO] Commento AI non disponibile per '{sezione}'.")
    return risposta


def _imposta_stili(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)


def _aggiungi_titolo(doc: Document, testo: str, livello: int = 1):
    heading = doc.add_heading(testo, level=livello)
    run = heading.runs[0]
    run.font.name = "Arial"
    if livello == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    elif livello == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    else:
        run.font.size = Pt(11)


def _aggiungi_tabella(doc: Document, df: pd.DataFrame):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Shading Accent 1"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_tuple in df.itertuples(index=False):
        row_cells = table.add_row().cells
        is_totale = str(row_tuple[0]).strip().lower() == "totale"
        for i, val in enumerate(row_tuple):
            col_name = str(df.columns[i]).lower()
            is_perc  = "%" in col_name or "peso" in col_name
            is_nodiv = any(k in col_name for k in ["dur.", "duration", "conv.", "convexity"])

            if isinstance(val, str):
                text = val
            elif isinstance(val, float):
                if is_perc:
                    text = f"{val:.2f}%"
                elif is_nodiv:
                    text = f"{val:.3f}"
                else:
                    text = f"{val:,.2f}"
            elif isinstance(val, int) and not isinstance(val, bool):
                text = f"{val:,}"
            else:
                text = str(val) if val is not None else ""

            cell = row_cells[i]
            cell.text = text
            p = cell.paragraphs[0]
            if isinstance(val, (int, float)):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.size = Pt(9)
            if is_totale:
                run.font.bold = True

    doc.add_paragraph()


def _ha_dati(dati: dict, chiave: str) -> bool:
    return (chiave in dati
            and dati[chiave] is not None
            and not (isinstance(dati[chiave], pd.DataFrame) and dati[chiave].empty))


def _fmt(val, is_perc=False) -> str:
    if val is None:
        return "n.d."
    return f"{val:.2f}%" if is_perc else f"{val:,.2f} €"


def genera_word(dati: dict, kpi: dict,
                nome_portafoglio: str = "Portafoglio Società/Gruppo",
                data_report: str = None,
                unita: str = "€") -> Document:

    doc = Document()
    _imposta_stili(doc)
    data_report = data_report or datetime.today().strftime("%d/%m/%Y")

    # Copertina
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nome_portafoglio.upper())
    run.font.bold  = True
    run.font.size  = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name  = "Arial"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"AR Investimenti\nReport del {data_report}")
    r2.font.size  = Pt(13)
    r2.font.name  = "Arial"
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # 1. EXECUTIVE SUMMARY
    # ------------------------------------------------------------------ #
    _aggiungi_titolo(doc, "1. Executive Summary", 1)
    kpi_text = json.dumps({
        "NAV totale":   _fmt(kpi.get("nav")),
        "Num. titoli":  kpi.get("n_titoli"),
        "P&L totale":   _fmt(kpi.get("pl_totale")),
        "Proventi":     _fmt(kpi.get("proventi")),
        "Rendimento %": _fmt(kpi.get("rendimento_%"), is_perc=True),
    }, ensure_ascii=False)

    commento = _commento_ai("Executive Summary", kpi_text, kpi)
    doc.add_paragraph(commento or (
        f"Il portafoglio presenta un valore complessivo di {_fmt(kpi.get('nav'))}, "
        f"composto da {kpi.get('n_titoli', 'n.d.')} titoli. "
        f"Il risultato economico complessivo ammonta a {_fmt(kpi.get('pl_totale'))}."
    ))

    # ------------------------------------------------------------------ #
    # 2. ANALISI PATRIMONIALE
    # ------------------------------------------------------------------ #
    _aggiungi_titolo(doc, "2. Analisi Patrimoniale", 1)

    sezioni_pat = [
        ("patrimoniale_asset_class",     "2.1 Composizione per Asset Class"),
        ("patrimoniale_fv_level",        "2.2 Distribuzione per Fair Value Level"),
        ("rating_governativi",           "2.3 Qualità creditizia – Titoli Governativi"),
        ("rating_non_governativi",       "2.4 Qualità creditizia – Titoli Non Governativi"),
        ("geografia_governativi",        "2.5 Distribuzione geografica – Governativi"),
        ("esposizione_valutaria",        "2.6 Esposizione Valutaria"),
        ("esposizione_settoriale",       "2.7 Esposizione Settoriale"),
        ("top_holdings",                 "2.8 Principali Holdings"),
        ("confronto_bv_fv",             "2.9 Book Value vs Fair Value"),
        ("composizione_valuation_class", "2.10 Composizione per Valuation Class"),
        ("oci_per_asset_class",         "2.11 OCI per Asset Class (IFRS9)"),
    ]

    for chiave, titolo_sez in sezioni_pat:
        if not _ha_dati(dati, chiave):
            continue
        _aggiungi_titolo(doc, titolo_sez, 2)
        _aggiungi_tabella(doc, dati[chiave])
        commento = _commento_ai(
            titolo_sez,
            dati[chiave].to_json(orient="records", force_ascii=False), kpi)
        if commento:
            doc.add_paragraph(commento)

    # ------------------------------------------------------------------ #
    # 3. ANALISI ECONOMICA
    # ------------------------------------------------------------------ #
    _aggiungi_titolo(doc, "3. Analisi Economica", 1)

    if _ha_dati(dati, "economica_completa"):
        _aggiungi_titolo(doc, "3.1 Risultato economico per Asset Class", 2)
        _aggiungi_tabella(doc, dati["economica_completa"])
        commento = _commento_ai(
            "3.1 Risultato economico per Asset Class",
            dati["economica_completa"].to_json(orient="records", force_ascii=False), kpi)
        if commento:
            doc.add_paragraph(commento)

    if _ha_dati(dati, "top_operazioni"):
        _aggiungi_titolo(doc, "3.2 Top 20 Operazioni di Periodo", 2)
        _aggiungi_tabella(doc, dati["top_operazioni"])
        commento = _commento_ai(
            "3.2 Top 20 Operazioni di Periodo",
            dati["top_operazioni"].to_json(orient="records", force_ascii=False), kpi)
        if commento:
            doc.add_paragraph(commento)

    # ------------------------------------------------------------------ #
    # 4. ANALISI EFFETTI
    # ------------------------------------------------------------------ #
    has_effetti = any(_ha_dati(dati, k)
                      for k in ["effetti_inventory", "effetti_tx_top20", "effetti_rie"])

    if has_effetti:
        _aggiungi_titolo(doc, "4. Analisi degli Effetti sul Portafoglio", 1)

        if _ha_dati(dati, "effetti_inventory"):
            _aggiungi_titolo(doc, "4.1 Effetto Nominale e di Mercato (Inventory N vs N-1)", 2)
            doc.add_paragraph(
                "La tabella seguente scompone la variazione del Fair Value tra i due "
                "esercizi in Effetto Nominale (variazione di quantità a prezzi N-1) "
                "e Effetto Mercato (rivalutazione della posizione finale). "
                "Il check è esatto per costruzione."
            )
            _aggiungi_tabella(doc, dati["effetti_inventory"])
            commento = _commento_ai(
                "4.1 Effetto Nominale e di Mercato",
                dati["effetti_inventory"].to_json(orient="records", force_ascii=False), kpi)
            if commento:
                doc.add_paragraph(commento)

        if _ha_dati(dati, "effetti_tx_top20"):
            _aggiungi_titolo(doc, "4.2 Top 20 Operazioni per Effetto Totale", 2)
            doc.add_paragraph(
                "Le 20 operazioni con maggiore impatto sul portafoglio, ordinate per "
                "valore assoluto dell'effetto totale (Effetto Nominale + Effetto Mercato)."
            )
            _aggiungi_tabella(doc, dati["effetti_tx_top20"])
            commento = _commento_ai(
                "4.2 Top 20 Operazioni per Effetto Totale",
                dati["effetti_tx_top20"].to_json(orient="records", force_ascii=False), kpi)
            if commento:
                doc.add_paragraph(commento)

        if _ha_dati(dati, "effetti_rie"):
            _aggiungi_titolo(doc, "4.3 Riepilogo Effetti per ISIN", 2)
            doc.add_paragraph(
                "Riepilogo per ISIN della decomposizione in tre effetti: "
                "Nominale, Prezzo (qualità di esecuzione rispetto al riferimento N-1) "
                "e Mercato (rivalutazione a fine periodo)."
            )
            _aggiungi_tabella(doc, dati["effetti_rie"])

    # ------------------------------------------------------------------ #
    # 5. ANALISI RISCHIO TASSO
    # ------------------------------------------------------------------ #
    has_rischio = any(_ha_dati(dati, k)
                      for k in ["scadenze_bucket", "duration_ponderata", "sensitivity_tassi"])

    if has_rischio:
        _aggiungi_titolo(doc, "5. Analisi Rischio Tasso", 1)

        if _ha_dati(dati, "scadenze_bucket"):
            _aggiungi_titolo(doc, "5.1 Distribuzione per Bucket di Scadenza", 2)
            _aggiungi_tabella(doc, dati["scadenze_bucket"])
            commento = _commento_ai(
                "Distribuzione scadenze",
                dati["scadenze_bucket"].to_json(orient="records", force_ascii=False), kpi)
            if commento:
                doc.add_paragraph(commento)

        if _ha_dati(dati, "duration_ponderata"):
            _aggiungi_titolo(doc, "5.2 Duration Ponderata per Asset Class", 2)
            _aggiungi_tabella(doc, dati["duration_ponderata"])
            commento = _commento_ai(
                "Duration ponderata",
                dati["duration_ponderata"].to_json(orient="records", force_ascii=False), kpi)
            if commento:
                doc.add_paragraph(commento)

        if _ha_dati(dati, "sensitivity_tassi"):
            _aggiungi_titolo(doc, "5.3 Stress Test Tassi – Approssimazione di Taylor", 2)
            doc.add_paragraph(
                "La tabella riporta la variazione stimata del valore di portafoglio "
                "per shift paralleli della curva dei tassi, calcolata con "
                "approssimazione di Taylor al secondo ordine: "
                "ΔP ≈ BV × (−D_mod × Δy + 1/2 × C × Δy²)."
            )
            df_sens = dati["sensitivity_tassi"]
            cols_word = [c for c in df_sens.columns if "(€)" not in str(c)]
            _aggiungi_tabella(doc, df_sens[cols_word])
            commento = _commento_ai(
                "Stress test tassi di interesse",
                dati["sensitivity_tassi"].to_json(orient="records", force_ascii=False), kpi)
            if commento:
                doc.add_paragraph(commento)

    # ------------------------------------------------------------------ #
    # 6. CONSIDERAZIONI FINALI
    # ------------------------------------------------------------------ #
    num_sez = 6 if has_rischio else (5 if has_effetti else 4)
    _aggiungi_titolo(doc, f"{num_sez}. Considerazioni Finali", 1)
    ctx = json.dumps({
        "kpi":             kpi_text,
        "analisi_prodotte": [k for k in dati if k not in ("kpi", "dettaglio")],
    }, ensure_ascii=False)
    conclusioni = _commento_ai(
        "Considerazioni finali e sintesi del portafoglio", ctx, kpi)
    doc.add_paragraph(conclusioni or "Inserire commento.")

    # Footer
    doc.add_paragraph()
    p_note = doc.add_paragraph(f"Report generato automaticamente il {data_report}.")
    p_note.runs[0].font.size = Pt(8)
    p_note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


def salva_word(doc: Document, path: str):
    doc.save(path)
    print(f"[OK] Word salvato: {path}")
